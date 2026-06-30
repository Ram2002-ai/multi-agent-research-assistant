"""Multi-format report export service."""

from __future__ import annotations

import csv
import io
import json
import re
import zipfile
from datetime import datetime, timezone
from html import escape
from pathlib import Path

from app.services.citations import CitationManager
from app.services.diagrams import mermaid_diagrams


class ExportUnavailable(RuntimeError):
    pass


class ExportService:
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.citations = CitationManager()

    def markdown(self, report: dict, sources: list[dict]) -> str:
        generated = report.get("completed_at") or report.get("created_at")
        bibliography = self.citations.bibliography_markdown(sources)
        diagrams = mermaid_diagrams(report["topic"])
        return (
            f"# Research Report: {report['topic']}\n\n"
            f"> Generated {generated} · Model `{report.get('model', 'default')}` · "
            f"{len(sources)} sources\n\n"
            "## Table of Contents\n\n"
            "1. [Research report](#research-report)\n"
            "2. [Pipeline diagram](#pipeline-diagram)\n"
            "3. [References](#references)\n\n"
            "## Research Report\n\n"
            f"{report.get('result', '')}\n\n"
            "## Pipeline Diagram\n\n"
            f"```mermaid\n{diagrams['pipeline']}\n```\n\n"
            "## References\n\n"
            f"{bibliography or '_No source URLs were returned by the model._'}\n"
        )

    def export(self, report: dict, sources: list[dict], format_name: str) -> Path:
        format_name = format_name.lower()
        safe_id = re.sub(r"[^a-zA-Z0-9_-]", "", report["id"])
        base = self.output_dir / f"report_{safe_id}"
        markdown = self.markdown(report, sources)
        if format_name in {"md", "markdown"}:
            path = base.with_suffix(".md")
            path.write_text(markdown, encoding="utf-8")
        elif format_name == "json":
            path = base.with_suffix(".json")
            path.write_text(
                json.dumps(
                    {"report": report, "sources": sources}, indent=2, ensure_ascii=False
                ),
                encoding="utf-8",
            )
        elif format_name == "html":
            path = base.with_suffix(".html")
            path.write_text(self._html(report, markdown), encoding="utf-8")
        elif format_name == "csv":
            path = base.with_name(base.name + "_citations").with_suffix(".csv")
            with path.open("w", newline="", encoding="utf-8-sig") as handle:
                writer = csv.DictWriter(
                    handle,
                    fieldnames=[
                        "title",
                        "url",
                        "domain",
                        "credibility_score",
                        "trust_score",
                        "authority_score",
                        "recency_score",
                        "citation_count",
                        "label",
                    ],
                    extrasaction="ignore",
                )
                writer.writeheader()
                writer.writerows(sources)
        elif format_name == "docx":
            path = base.with_suffix(".docx")
            self._docx(path, report, sources)
        elif format_name == "pdf":
            path = base.with_suffix(".pdf")
            self._pdf(path, report)
        elif format_name == "zip":
            path = base.with_suffix(".zip")
            self._zip(path, report, sources)
        else:
            raise ExportUnavailable(f"Unsupported export format: {format_name}")
        return path

    def _html(self, report: dict, markdown: str) -> str:
        try:
            import markdown as md

            body = md.markdown(markdown, extensions=["fenced_code", "tables", "toc"])
        except ImportError:
            body = f"<pre>{escape(markdown)}</pre>"
        return f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width">
<title>{escape(report['topic'])} · ResearchOS</title>
<style>body{{font:16px/1.65 Inter,system-ui;max-width:900px;margin:60px auto;padding:0 24px;color:#17211b}}
h1,h2{{color:#163f2d}}blockquote{{border-left:4px solid #2d7a52;padding:12px 20px;background:#f0f7f3}}
code,pre{{background:#f4f4f0;padding:3px 6px;border-radius:6px}}pre{{padding:18px;overflow:auto}}</style>
</head><body>{body}</body></html>"""

    def _docx(self, path: Path, report: dict, sources: list[dict]) -> None:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise ExportUnavailable("DOCX export requires python-docx") from exc
        document = Document()
        title = document.add_heading("ResearchOS", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = document.add_paragraph(report["topic"])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(
            f"Generated {datetime.now(timezone.utc).strftime('%B %d, %Y')}"
        )
        document.add_page_break()
        document.add_heading("Research Report", level=1)
        for block in report.get("result", "").split("\n\n"):
            if block.startswith("## "):
                document.add_heading(block[3:], level=2)
            else:
                document.add_paragraph(block)
        document.add_heading("References", level=1)
        for citation in self.citations.format_all(sources)["ieee"]:
            document.add_paragraph(citation)
        document.save(path)

    def _pdf(self, path: Path, report: dict) -> None:
        try:
            from pdf_formatter import generate_pdf
        except ImportError as exc:
            raise ExportUnavailable("PDF export requires reportlab") from exc
        generate_pdf(report["topic"], report.get("result", ""), str(path))

    def _zip(self, path: Path, report: dict, sources: list[dict]) -> None:
        markdown = self.markdown(report, sources)
        citations = io.StringIO()
        writer = csv.DictWriter(
            citations,
            fieldnames=["title", "url", "domain", "credibility_score", "label"],
            extrasaction="ignore",
        )
        writer.writeheader()
        writer.writerows(sources)
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("report.md", markdown)
            archive.writestr(
                "report.json",
                json.dumps({"report": report, "sources": sources}, indent=2),
            )
            archive.writestr("citations.csv", citations.getvalue())
            archive.writestr("report.html", self._html(report, markdown))
